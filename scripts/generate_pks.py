#!/usr/bin/env python3
"""
Generate a PKS (Perjanjian Kerja Sama) DOCX by filling the official UPI
template's placeholders — mirrors the manual workflow of replacing the
yellow-highlighted fields in "Cth. PKS atau MoU.docx".

The legal text (all pasal, the PIHAK KESATU / UPI side) stays verbatim;
only the partner-specific {{TOKEN}} placeholders are filled.

Usage:
  python generate_pks.py --data data.json --output PKS_Company.docx [--template tpl.docx]
"""

import argparse
import json
import sys
from pathlib import Path


TOOLS_DIR     = Path.home() / ".claude" / "magang-tools"
CONFIG_PATH   = TOOLS_DIR / "config.json"
BUNDLED_TPL   = TOOLS_DIR / "templates" / "pks_template.docx"

# data key -> template token
TOKENS = {
    "partner_org_upper":  "PARTNER_ORG_UPPER",
    "partner_org_sub":    "PARTNER_ORG_SUB",
    "partner_org":        "PARTNER_ORG",
    "partner_rep_name":   "PARTNER_REP_NAME",
    "partner_rep_title":  "PARTNER_REP_TITLE",
    "partner_address":    "PARTNER_ADDRESS",
    "partner_phone":      "PARTNER_PHONE",
    "partner_fax":        "PARTNER_FAX",
    "partner_email":      "PARTNER_EMAIL",
    "partner_region":     "PARTNER_REGION",
    "date_hari":          "DATE_HARI",
    "date_tgl_word":      "DATE_TGL_WORD",
    "date_bulan":         "DATE_BULAN",
    "date_tahun_word":    "DATE_TAHUN_WORD",
    "date_numeric":       "DATE_NUMERIC",
}


def load_config() -> dict:
    if CONFIG_PATH.exists():
        return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    return {}


def versioned_path(base: Path) -> Path:
    if not base.exists():
        return base
    v = 2
    while True:
        candidate = base.parent / f"{base.stem}_v{v}{base.suffix}"
        if not candidate.exists():
            return candidate
        v += 1


def resolve_template(cli_template: str, config: dict) -> Path:
    if cli_template:
        return Path(cli_template).expanduser()
    cfg = config.get("pks_template_path")
    if cfg:
        return Path(cfg).expanduser()
    return BUNDLED_TPL


def fill_runs(paragraphs, repl: dict) -> int:
    """Replace {{TOKEN}} substrings in each run; tokens are single runs in
    the tokenized template, so substring replace is exact and safe."""
    n = 0
    for para in paragraphs:
        for run in para.runs:
            if "{{" not in run.text:
                continue
            new = run.text
            for token, value in repl.items():
                if token in new:
                    new = new.replace(token, value)
                    n += 1
            if new != run.text:
                run.text = new
    return n


def fill_template(template: Path, data: dict, output_path: Path) -> Path:
    from docx import Document

    if not template.exists():
        print(json.dumps({"error": f"Template not found: {template}"}))
        sys.exit(3)

    doc = Document(str(template))

    # Build {{TOKEN}} -> value map. Missing values default to empty so the
    # placeholder disappears rather than leaking into the document.
    repl = {}
    for key, token in TOKENS.items():
        repl["{{" + token + "}}"] = str(data.get(key, "") or "")

    fill_runs(doc.paragraphs, repl)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                fill_runs(cell.paragraphs, repl)

    output_path = versioned_path(output_path)
    doc.save(str(output_path))
    return output_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data",     required=True)
    parser.add_argument("--output",   required=True)
    parser.add_argument("--template", default=None)
    args = parser.parse_args()

    data     = json.loads(Path(args.data).read_text(encoding="utf-8"))
    config   = load_config()
    template = resolve_template(args.template, config)
    out      = fill_template(template, data, Path(args.output).expanduser())
    print(json.dumps({"success": True, "output": str(out),
                      "template": str(template)}))


if __name__ == "__main__":
    main()
