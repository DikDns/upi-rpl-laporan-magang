#!/usr/bin/env python3
"""
Generate the official "Lembar Penilaian Penyelia" DOCX (pedoman v6) from JSON.
Blank form by default; fills scores + computes Jumlah/Rata-rata when `nilai`
is supplied. Mirrors generate_logbook.py conventions.

Usage: python generate_penilaian.py --data data.json --output out.docx
"""

import argparse
import json
import sys
from pathlib import Path


TOOLS_DIR = Path.home() / ".claude" / "magang-tools"
CONFIG_PATH = TOOLS_DIR / "config.json"

PENILAIAN_FALLBACK = {
    "title": "PENILAIAN PELAKSANAAN KEGIATAN MBKM PROGRAM MSIB / P3NK (MAGANG MANDIRI)",
    "scale": "0 – 100",
    "note": "Jika sudah ada form dari instansi, form penilaian ini tidak berlaku.",
    "indicators": [
        "Pemahaman terhadap tugas dan tanggung jawab",
        "Penguasaan konsep dan teori yang relevan",
        "Kemampuan dalam pemrograman / teknis",
        "Penggunaan teknologi, tools, dan framework",
        "Kualitas hasil kerja",
        "Kemampuan analisis, pemecahan masalah & inisiatif",
        "Komunikasi dan kolaborasi dalam tim",
        "Manajemen waktu dan ketepatan penyelesaian tugas",
        "Adaptasi terhadap lingkungan kerja dan perubahan",
        "Etika, sikap profesional, dan tanggung jawab",
    ],
}

INTRO = (
    "Dengan mempertimbangkan segala aspek, baik dari segi bobot pekerjaan maupun "
    "pelaksanaan kegiatan MBKM Program MSIB / P3NK (Magang Mandiri), maka kami "
    "memutuskan bahwa yang bersangkutan telah menyelesaikan kewajibannya, "
    "sebagaimana terinci dalam indikator penilaian sebagai berikut:"
)


def load_config() -> dict:
    if CONFIG_PATH.exists():
        return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    return {}


def versioned_path(base: Path) -> Path:
    if not base.exists():
        return base
    v = 2
    while True:
        candidate = base.parent / f"{base.stem}_v{v}{base.suffix}"
        if not candidate.exists():
            return candidate
        v += 1


def set_cell_border(cell):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_fixed_layout(table, widths):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.tblPr
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for el in tbl.findall(qn("w:tblGrid")):
        tbl.remove(el)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w.twips)))
        grid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, grid)
    table.allow_autofit = False


def generate(data: dict, config: dict, output_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pen = config.get("penilaian_penyelia") or PENILAIAN_FALLBACK
    indicators = pen.get("indicators") or PENILAIAN_FALLBACK["indicators"]
    title = pen.get("title", PENILAIAN_FALLBACK["title"])
    scale = pen.get("scale", "0 – 100")
    note = pen.get("note", PENILAIAN_FALLBACK["note"])

    fmt = config.get("formatting", {})
    margins = fmt.get("margins_cm", {"left": 4, "top": 3, "bottom": 3, "right": 3})
    font_name = fmt.get("font", "Arial")
    font_size = fmt.get("font_size_body", 11)

    doc = Document()
    sec = doc.sections[0]
    _pg = {"a4": (21.0, 29.7), "letter": (21.59, 27.94)}
    _pw, _ph = _pg.get(str(fmt.get("page_size", "A4")).lower(), (21.0, 29.7))
    sec.page_width = Cm(_pw); sec.page_height = Cm(_ph)
    sec.left_margin = Cm(margins["left"]); sec.right_margin = Cm(margins["right"])
    sec.top_margin = Cm(margins["top"]); sec.bottom_margin = Cm(margins["bottom"])

    printable = _pw - margins["left"] - margins["right"]  # cm

    def run_fmt(run, size=None, bold=False, italic=False, color=None):
        run.bold = bold; run.italic = italic
        run.font.name = font_name
        run.font.size = Pt(size or font_size)
        if color is not None:
            run.font.color.rgb = color

    def add_para(text, align=None, size=None, bold=False, italic=False,
                 color=None, space_after=2):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        run_fmt(p.add_run(text), size, bold, italic, color)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    # 1) KOP placeholder
    add_para("[ KOP SURAT INSTITUSI MITRA ]", align=WD_ALIGN_PARAGRAPH.CENTER,
             italic=True, color=RGBColor(0x80, 0x80, 0x80))
    doc.add_paragraph()

    # 2) Title
    add_para(title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_paragraph()

    # 3) Identity block (label | : | value), borderless
    fields = [
        ("Nama Penyelia", data.get("nama_penyelia", "")),
        ("Nama Institusi Mitra", data.get("nama_institusi", "")),
        ("Nama Mahasiswa", data.get("nama_mahasiswa", "")),
        ("Nomor Induk Mahasiswa", data.get("nim", "")),
        ("Waktu Pelaksanaan", data.get("waktu_pelaksanaan", "")),
    ]
    wlab, wcol = Cm(5.0), Cm(0.4)
    wval = Cm(max(printable - 5.0 - 0.4, 4.0))
    idt = doc.add_table(rows=len(fields), cols=3)
    set_fixed_layout(idt, [wlab, wcol, wval])
    for i, (label, value) in enumerate(fields):
        cells = idt.rows[i].cells
        for cell, w, txt in ((cells[0], wlab, label), (cells[1], wcol, ":"), (cells[2], wval, value)):
            cell.width = w
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            run_fmt(p.add_run(txt))
    doc.add_paragraph()

    # 4) Intro
    add_para(INTRO, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
    doc.add_paragraph()

    # 5) Indicator table
    nilai = data.get("nilai") or {}
    # normalize keys to 1-based string indices
    nilai = {str(k): v for k, v in nilai.items()}
    w_no, w_nilai = 1.2, 2.2
    w_ind = max(printable - w_no - w_nilai, 5.0)
    widths = [Cm(w_no), Cm(w_ind), Cm(w_nilai)]
    n_rows = 1 + len(indicators) + 2  # header + indicators + Jumlah + Rata-rata
    tbl = doc.add_table(rows=n_rows, cols=3)
    tbl.style = "Table Grid"
    set_fixed_layout(tbl, widths)

    def cell_text(cell, text, w, align=None, bold=False):
        cell.width = w
        set_cell_border(cell)
        p = cell.paragraphs[0]
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        run_fmt(p.add_run(text), bold=bold)

    # header
    hdr = tbl.rows[0].cells
    cell_text(hdr[0], "No.", widths[0], WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    cell_text(hdr[1], "Indikator Penilaian", widths[1], WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    cell_text(hdr[2], "Nilai*)", widths[2], WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    total = 0.0
    have_any = False
    for idx, ind in enumerate(indicators):
        row = tbl.rows[idx + 1].cells
        cell_text(row[0], str(idx + 1), widths[0], WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row[1], ind, widths[1])
        val = nilai.get(str(idx + 1), "")
        val_txt = ""
        if val != "" and val is not None:
            try:
                num = float(val)
                total += num; have_any = True
                val_txt = str(int(num)) if num == int(num) else str(num)
            except (TypeError, ValueError):
                val_txt = str(val)
        cell_text(row[2], val_txt, widths[2], WD_ALIGN_PARAGRAPH.CENTER)

    # Jumlah + Rata-rata (computed only when scores supplied; divisor fixed at len(indicators))
    jumlah_txt = rata_txt = ""
    if have_any:
        avg = total / len(indicators)
        jumlah_txt = str(int(total)) if total == int(total) else str(round(total, 2))
        rata_txt = str(round(avg, 2))
    j = tbl.rows[1 + len(indicators)].cells
    cell_text(j[0], "", widths[0])
    cell_text(j[1], "Jumlah", widths[1], bold=True)
    cell_text(j[2], jumlah_txt, widths[2], WD_ALIGN_PARAGRAPH.CENTER)
    r = tbl.rows[2 + len(indicators)].cells
    cell_text(r[0], "", widths[0])
    cell_text(r[1], "Rata-rata", widths[1], bold=True)
    cell_text(r[2], rata_txt, widths[2], WD_ALIGN_PARAGRAPH.CENTER)

    # 6) Footnotes
    add_para(f"*) Skala/Rentang penilaian Penyelia {scale}.", size=font_size - 2, space_after=0)
    add_para(f"**) {note}", size=font_size - 2)

    # 7) Signature
    doc.add_paragraph()
    add_para(data.get("tempat_tanggal", ""), align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_para("Penyelia,", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    sig = doc.add_table(rows=1, cols=1)
    tr = sig.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), "1700"); trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)
    name_line = data.get("nama_penyelia", "") or "______________________________________"
    add_para(name_line, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, space_after=0)
    add_para("Tandatangan, Nama, Posisi, dan di Stempel",
             align=WD_ALIGN_PARAGRAPH.RIGHT, size=font_size - 2)

    output_path = versioned_path(output_path)
    doc.save(str(output_path))
    return output_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    data = json.loads(Path(args.data).read_text(encoding="utf-8"))
    config = load_config()
    out = generate(data, config, Path(args.output).expanduser())
    print(json.dumps({"success": True, "output": str(out)}))


if __name__ == "__main__":
    main()
