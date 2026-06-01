#!/usr/bin/env python3
"""
Compile laporan section .md files into a single A4 DOCX.

Handles, per UPI/campus convention (sizes/margins/font from config.json):
  - cover & lembar-pengesahan: full-page templates rendered from key:value
    data (not markdown), spaced to fill the page
  - markdown bab sections: campus headings (BAB 14pt center, sub-bab 12pt
    left, black), lists, pipe tables, **bold**/*italic*, and embedded
    images via ![caption](path) with auto-numbered "Gambar X.Y" captions
Sections are ordered by SECTIONS_ORDER, one page break between each.

Usage:
  python generate_laporan.py --sections-dir ./laporan-draft --output Laporan.docx
"""

import argparse
import json
import re
import sys
from pathlib import Path


TOOLS_DIR   = Path.home() / ".claude" / "magang-tools"
CONFIG_PATH = TOOLS_DIR / "config.json"

SECTIONS_ORDER = [
    "cover",
    "lembar-pengesahan",
    "kata-pengantar",
    "daftar-isi",
    "bab1",
    "bab2",
    "bab3",
    "bab4",
    "daftar-pustaka",
    "lampiran",
]


def load_config() -> dict:
    if CONFIG_PATH.exists():
        return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    return {}


def versioned_path(base: Path) -> Path:
    if not base.exists():
        return base
    v = 2
    while True:
        candidate = base.parent / f"{base.stem}_v{v}{base.suffix}"
        if not candidate.exists():
            return candidate
        v += 1


def apply_inline(paragraph, text: str, font_name: str, font_size_pt):
    from docx.shared import Pt

    # Split on **bold**, *italic*, and `code`
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
        elif len(part) >= 2 and part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(font_size_pt - 1)
        else:
            run = paragraph.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)


def set_fixed_layout(table, widths):
    """Force fixed table layout so per-cell widths are honored by Word.
    `widths` is a list of docx Length (e.g. Cm); their sum should fit the
    printable area or Word will re-distribute regardless of fixed layout."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for el in tbl.findall(qn("w:tblGrid")):
        tbl.remove(el)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w.twips)))
        grid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, grid)
    table.allow_autofit = False


def md_to_doc(doc, md_text: str, font_name: str, font_size: int,
              base_dir: Path = None, chapter_label: str = None,
              fig_counter: dict = None, heading_sizes: dict = None):
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if fig_counter is None:
        fig_counter = {}

    # Campus heading convention (from reference proposal), NOT Word's blue
    # built-in heading styles: black, bold, document font, explicit sizes.
    #   H1 (BAB)  -> chapter size, CENTER
    #   H2 (x.y)  -> section size, LEFT
    #   H3 (x.y.z)-> section size, LEFT
    #   H4        -> body size, LEFT
    hs = heading_sizes or {
        1: (14, "center"), 2: (12, "left"),
        3: (12, "left"),   4: (font_size, "left"),
    }
    BLACK = RGBColor(0, 0, 0)

    def add_heading(level, text):
        size, align = hs.get(level, (font_size, "left"))
        p = doc.add_paragraph()
        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align == "center"
                       else WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.name = font_name
        r.font.size = Pt(size)
        r.font.color.rgb = BLACK
        return p

    # Usable width on A4 with 4cm/3cm margins ≈ 14cm.
    IMG_WIDTH = Cm(14)

    def add_image(alt_text: str, img_ref: str):
        # Resolve path relative to the section's directory.
        p = Path(img_ref).expanduser()
        if not p.is_absolute() and base_dir is not None:
            p = (base_dir / img_ref).resolve()
        # Figure number scoped to chapter label (e.g. "3" -> Gambar 3.1).
        key = chapter_label or "_"
        fig_counter[key] = fig_counter.get(key, 0) + 1
        n = fig_counter[key]
        caption = (f"Gambar {chapter_label}.{n}" if chapter_label
                   else f"Gambar {n}")
        if alt_text.strip():
            caption += f" {alt_text.strip()}"

        if not p.exists():
            warn = doc.add_paragraph()
            warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = warn.add_run(f"[Gambar tidak ditemukan: {img_ref}]")
            r.italic = True
            r.font.name = font_name
            r.font.size = Pt(font_size)
        else:
            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_p.add_run()
            run.add_picture(str(p), width=IMG_WIDTH)

        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap_p.add_run(caption)
        cr.font.name = font_name
        cr.font.size = Pt(font_size - 1)
        cap_p.paragraph_format.space_after = Pt(6)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # ── Image: ![caption](path) on its own line ──
        img_m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if img_m:
            add_image(img_m.group(1), img_m.group(2).strip())
            i += 1
            continue

        # ── Headings (longest prefix first) ──
        if line.startswith("#### "):
            add_heading(4, line[5:].strip())
            i += 1
            continue
        if line.startswith("### "):
            add_heading(3, line[4:].strip())
            i += 1
            continue
        if line.startswith("## "):
            add_heading(2, line[3:].strip())
            i += 1
            continue
        if line.startswith("# "):
            add_heading(1, line[2:].strip())
            i += 1
            continue

        # ── Bullet list ──
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            apply_inline(p, line[2:].strip(), font_name, font_size)
            i += 1
            continue

        # ── Numbered list ──
        num_m = re.match(r"^(\d+)\.\s+(.+)", line)
        if num_m:
            p = doc.add_paragraph(style="List Number")
            apply_inline(p, num_m.group(2).strip(), font_name, font_size)
            i += 1
            continue

        # ── Table (markdown pipe table) ──
        if "|" in line:
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            # Filter out separator rows (---|---)
            data_rows = [r for r in table_lines
                         if not re.match(r"^\s*[\|:\-\s]+$", r)]
            if data_rows:
                parsed = []
                for row in data_rows:
                    cells = [c.strip() for c in row.strip().strip("|").split("|")]
                    parsed.append(cells)
                if parsed:
                    from docx.shared import Emu
                    max_cols = max(len(r) for r in parsed)
                    tbl = doc.add_table(rows=len(parsed), cols=max_cols)
                    tbl.style = "Table Grid"
                    # Fit columns within printable width so Word does not
                    # re-distribute / overflow the page (equal columns).
                    sec0 = doc.sections[0]
                    usable = sec0.page_width - sec0.left_margin - sec0.right_margin
                    col_w = Emu(int(usable) // max_cols)
                    set_fixed_layout(tbl, [col_w] * max_cols)
                    for ri, row in enumerate(parsed):
                        for ci, cell_text in enumerate(row):
                            if ci < max_cols:
                                cell = tbl.cell(ri, ci)
                                cell.width = col_w
                                p = cell.paragraphs[0]
                                apply_inline(p, cell_text, font_name, font_size)
                                if ri == 0:
                                    for r in p.runs:
                                        r.bold = True
            continue

        # ── Horizontal rule ──
        if re.match(r"^[-*_]{3,}$", line.strip()):
            doc.add_paragraph()
            i += 1
            continue

        # ── Regular paragraph ──
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(0)
        apply_inline(p, line.strip(), font_name, font_size)
        i += 1


ASSETS_DIR = TOOLS_DIR / "assets"


def parse_kv(md_text: str) -> dict:
    """Parse simple `key: value` lines (cover / lembar-pengesahan data)."""
    out = {}
    for line in md_text.splitlines():
        m = re.match(r"^\s*([A-Za-z_]+)\s*:\s*(.*)$", line)
        if m:
            out[m.group(1).strip().lower()] = m.group(2).strip()
    return out


def scan_has_tables(sections_dir: Path) -> bool:
    """Return True if any bab .md file in sections_dir contains a markdown table."""
    for md in sections_dir.glob("bab*.md"):
        if "|" in md.read_text(encoding="utf-8"):
            return True
    return False


def scan_has_images(sections_dir: Path) -> bool:
    """Return True if any bab .md file in sections_dir contains an image reference."""
    for md in sections_dir.glob("bab*.md"):
        if re.search(r"!\[", md.read_text(encoding="utf-8")):
            return True
    return False


def _insert_toc_field(doc, font_name: str, font_size: int, title: str, toc_instruction: str):
    """Insert a section title (H1 style) and a Word TOC field paragraph."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    title_p.paragraph_format.line_spacing = 1.0
    r = title_p.add_run(title)
    r.bold = True
    r.font.name = font_name
    r.font.size = Pt(14)

    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.space_after = Pt(0)

    run1 = toc_p.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    run1._r.append(fc1)

    run2 = toc_p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {toc_instruction} "
    run2._r.append(instr)

    run3 = toc_p.add_run()
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    run3._r.append(fc3)


def _gap(doc, n=1):
    """n empty single-spaced lines (~0.39cm each) for vertical spacing.
    Explicit spacing is reliable across viewers; Word's vertical-justify
    (w:vAlign) is ignored by Preview/Quick Look/LibreOffice."""
    from docx.shared import Pt
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def _block(doc, font_name, lines, align="center"):
    """One paragraph holding several lines (line breaks within), so the
    block stays tight while Word's vertical-justify spreads the BLOCKS
    across the page. `lines` = list of (text, size_pt, bold)."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align == "center"
                   else WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    for idx, (text, size, bold) in enumerate(lines):
        r = p.add_run(text)
        r.bold = bold
        r.font.name = font_name
        r.font.size = Pt(size)
        if idx < len(lines) - 1:
            r.add_break()
    return p


def render_cover(doc, fields: dict, config: dict):
    """Full-page cover per pedoman (Lampiran Contoh Cover, hal. 13).

    Emits five tight blocks; the section's vertical-justify spreads them
    from top to bottom so the cover fills the whole page."""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    fmt       = config.get("formatting", {})
    font_name = fmt.get("font", "Arial")
    program   = config.get("program", "Rekayasa Perangkat Lunak")
    campus    = config.get("campus", "Kampus UPI di Cibiru").upper()

    nama  = fields.get("nama", "Nama Mahasiswa")
    nim   = fields.get("nim", "NIM")
    tahun = fields.get("tahun", "Tahun")

    # Gaps tuned to spread 5 blocks across A4 usable height (~23.7cm):
    # title near top, logo + identitas mid, institusi near bottom.
    _block(doc, font_name, [
        ("Laporan Pelaksanaan Kegiatan MBKM", 12, True),
        ("MBKM Program MSIB / P3NK (Magang Mandiri)", 12, True),
    ])
    _gap(doc, 1)
    _block(doc, font_name, [
        ("Diajukan sebagai salah satu syarat Kegiatan MBKM", 10, False),
        (f"pada Program Studi {program}", 10, False),
    ])
    _gap(doc, 8)

    # LOGO UPI — embed asset if present, else placeholder text.
    logo = ASSETS_DIR / "upi-logo.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.add_run().add_picture(str(logo), height=Cm(4))
    else:
        _block(doc, font_name, [("LOGO UPI", 12, False)])

    _gap(doc, 8)
    _block(doc, font_name, [
        ("Oleh.", 12, False),
        (nama, 12, True),
        (nim, 12, False),
    ])
    _gap(doc, 13)
    _block(doc, font_name, [
        (campus, 14, True),
        (program.upper(), 14, True),
        ("UNIVERSITAS PENDIDIKAN INDONESIA", 14, True),
        (str(tahun), 14, True),
    ])


def render_lembar_pengesahan(doc, fields: dict, config: dict):
    """Lembar Pengesahan per pedoman (Lampiran Contoh, hal. 14).

    Tight blocks + vertical-justify so the signature block reaches the
    lower part of the page like the pedoman example."""
    from docx.shared import Pt

    fmt       = config.get("formatting", {})
    font_name = fmt.get("font", "Arial")
    font_size = fmt.get("font_size_body", 11)
    program   = config.get("program", "Rekayasa Perangkat Lunak")

    dosen   = fields.get("dosen_pembimbing", "")
    penyelia = fields.get("penyelia", "")
    kaprodi  = fields.get("kaprodi", "")
    kaprodi_nip = fields.get("kaprodi_nip", "")

    _block(doc, font_name, [
        ("Laporan Pelaksanaan Kegiatan MBKM", 14, True),
        ("MBKM Program MSIB / P3NK (Magang Mandiri)", 14, True),
    ])
    _gap(doc, 4)
    _block(doc, font_name, [("Lembar Pengesahan", 12, True)])
    _gap(doc, 5)
    _block(doc, font_name, [
        ("Diajukan sebagai salah satu syarat kegiatan MBKM", 10, False),
        (f"pada Program Studi {program}", 10, False),
    ])
    _gap(doc, 8)
    _block(doc, font_name, [
        (f"{'Dosen Pembimbing':<18}: {dosen}", font_size, False),
        (f"{'Penyelia':<18}: {penyelia}", font_size, False),
    ], align="left")
    _gap(doc, 10)
    _block(doc, font_name, [
        ("Mengetahui,", font_size, False),
        (f"Ketua Program Studi {program},", font_size, False),
    ])
    _gap(doc, 6)
    sig_lines = [("__________________________", font_size, False),
                 (kaprodi if kaprodi else "Tandatangan, Nama Jelas & NIP",
                  font_size, bool(kaprodi))]
    if kaprodi_nip:
        sig_lines.append((f"NIP. {kaprodi_nip}", font_size, False))
    _block(doc, font_name, sig_lines)


def compile_sections(sections_dir: Path, output_path: Path, config: dict) -> Path:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    fmt       = config.get("formatting", {})
    margins   = fmt.get("margins_cm", {"left": 4, "top": 3, "bottom": 3, "right": 3})
    font_name = fmt.get("font", "Arial")
    font_size = fmt.get("font_size_body", 11)
    spacing   = fmt.get("line_spacing", 1.5)
    h_chapter = fmt.get("font_size_chapter_title", 14)
    h_section = fmt.get("font_size_section_title", 12)
    heading_sizes = {
        1: (h_chapter, "center"),
        2: (h_section, "left"),
        3: (h_section, "left"),
        4: (font_size, "left"),
    }

    # Page size: pedoman specifies only margins, so default A4 (Indonesian
    # academic standard); overridable via config formatting.page_size.
    PAGE = {"a4": (21.0, 29.7), "letter": (21.59, 27.94)}
    pw, ph = PAGE.get(str(fmt.get("page_size", "A4")).lower(), (21.0, 29.7))

    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Cm(pw)
    sec.page_height   = Cm(ph)
    sec.left_margin   = Cm(margins["left"])
    sec.right_margin  = Cm(margins["right"])
    sec.top_margin    = Cm(margins["top"])
    sec.bottom_margin = Cm(margins["bottom"])

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    # line_spacing as a float → WD_LINE_SPACING.MULTIPLE (e.g. 1.5x).
    # Passing a Length here would switch to EXACTLY mode and overflow.
    style.paragraph_format.line_spacing = spacing

    # Collect sections in order
    all_md = {f.stem: f for f in sections_dir.glob("*.md")}
    ordered = []
    for name in SECTIONS_ORDER:
        if name == "daftar-isi":
            ordered.append((name, None))   # auto-generated, no .md needed
        elif name in all_md:
            ordered.append((name, all_md.pop(name)))
    for name, path in sorted(all_md.items()):
        ordered.append((name, path))

    # Scan for conditional daftar sections
    has_tables = scan_has_tables(sections_dir)
    has_images = scan_has_images(sections_dir)

    # Figure numbering is shared across the whole document, scoped per
    # chapter label (bab3 -> "Gambar 3.x", lampiran -> "Gambar L.x").
    fig_counter = {}

    def chapter_label_for(name: str):
        m = re.search(r"bab(\d+)", name)
        if m:
            return m.group(1)
        if "lampiran" in name:
            return "L"
        return None

    for idx, (name, md_file) in enumerate(ordered):
        if idx > 0:
            doc.add_page_break()
        md_content = md_file.read_text(encoding="utf-8") if md_file else ""
        # Cover and Lembar Pengesahan are fixed full-page templates with
        # per-line font sizes — rendered from key:value data, not markdown.
        # They distribute their own vertical spacing to fill the A4 page.
        if name == "cover":
            render_cover(doc, parse_kv(md_content), config)
        elif name == "lembar-pengesahan":
            render_lembar_pengesahan(doc, parse_kv(md_content), config)
        elif name == "daftar-isi":
            _insert_toc_field(doc, font_name, font_size, "DAFTAR ISI",
                              'TOC \\o "1-3" \\h \\z \\u')
            if has_tables:
                doc.add_page_break()
                _insert_toc_field(doc, font_name, font_size, "DAFTAR TABEL",
                                  'TOC \\h \\z \\c "Tabel"')
            if has_images:
                doc.add_page_break()
                _insert_toc_field(doc, font_name, font_size, "DAFTAR GAMBAR",
                                  'TOC \\h \\z \\c "Gambar"')
        else:
            md_to_doc(doc, md_content, font_name, font_size,
                      base_dir=sections_dir,
                      chapter_label=chapter_label_for(name),
                      fig_counter=fig_counter,
                      heading_sizes=heading_sizes)

    output_path = versioned_path(output_path)
    doc.save(str(output_path))
    return output_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--sections-dir", required=True)
    parser.add_argument("--output",       required=True)
    args = parser.parse_args()

    sections_dir = Path(args.sections_dir).expanduser().resolve()
    if not sections_dir.exists():
        print(json.dumps({"error": f"Not found: {sections_dir}"}))
        sys.exit(1)

    config     = load_config()
    out        = compile_sections(sections_dir, Path(args.output).expanduser(), config)
    print(json.dumps({"success": True, "output": str(out)}))


if __name__ == "__main__":
    main()
