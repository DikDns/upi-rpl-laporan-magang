#!/usr/bin/env python3
"""
Generate logbook DOCX from JSON data.
Usage: python generate_logbook.py --data /path/to/data.json --output /path/to/output.docx
"""

import argparse
import json
import sys
from pathlib import Path


TOOLS_DIR = Path.home() / ".claude" / "magang-tools"
CONFIG_PATH = TOOLS_DIR / "config.json"


def load_config() -> dict:
    if CONFIG_PATH.exists():
        return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    return {}


def versioned_path(base: Path) -> Path:
    if not base.exists():
        return base
    v = 2
    while True:
        candidate = base.parent / f"{base.stem}_v{v}{base.suffix}"
        if not candidate.exists():
            return candidate
        v += 1


def set_cell_border(cell):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_fixed_layout(table, widths):
    """Force fixed table layout so per-cell widths are honored by Word."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr
    # remove existing layout if any
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # explicit grid columns (EMU -> twips)
    for el in tbl.findall(qn("w:tblGrid")):
        tbl.remove(el)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w.twips)))
        grid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, grid)
    table.allow_autofit = False


def generate(data: dict, config: dict, output_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    fmt = config.get("formatting", {})
    margins = fmt.get("margins_cm", {"left": 4, "top": 3, "bottom": 3, "right": 3})
    font_name = fmt.get("font", "Arial")
    font_size = fmt.get("font_size_body", 11)

    sec = doc.sections[0]
    # A4 (Indonesian academic standard); pedoman specifies only margins.
    _pg = {"a4": (21.0, 29.7), "letter": (21.59, 27.94)}
    _pw, _ph = _pg.get(str(fmt.get("page_size", "A4")).lower(), (21.0, 29.7))
    sec.page_width    = Cm(_pw)
    sec.page_height   = Cm(_ph)
    sec.left_margin   = Cm(margins["left"])
    sec.right_margin  = Cm(margins["right"])
    sec.top_margin    = Cm(margins["top"])
    sec.bottom_margin = Cm(margins["bottom"])

    def add_centered_bold(text, size=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.name = font_name
        run.font.size = Pt(size or font_size)
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_identity(fields):
        """Borderless 3-col table (label | : | value) so colons align."""
        wlab, wcol, wval = Cm(4.2), Cm(0.4), Cm(9.4)
        t = doc.add_table(rows=len(fields), cols=3)
        set_fixed_layout(t, [wlab, wcol, wval])
        for i, (label, value) in enumerate(fields):
            cells = t.rows[i].cells
            for cell, w, txt in ((cells[0], wlab, label), (cells[1], wcol, ":"), (cells[2], wval, value)):
                cell.width = w
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(txt)
                run.font.name = font_name
                run.font.size = Pt(font_size)

    # ── header ──
    tmpl = config.get("logbook_template", {})
    add_centered_bold(tmpl.get("title", "CATATAN HARIAN & KEHADIRAN PESERTA"))
    add_centered_bold(tmpl.get("subtitle",
        "KEGIATAN MBKM KEGIATAN PROGRAM MSIB / P3NK (MAGANG MANDIRI) PADA MITRA"))
    doc.add_paragraph()

    add_identity([
        ("Nama Mahasiswa", data.get("nama_mahasiswa", "")),
        ("NIM",            data.get("nim", "")),
        ("Nama Mitra",     data.get("nama_mitra", "")),
        ("Nama Penyelia",  data.get("nama_penyelia", "")),
    ])
    doc.add_paragraph()

    # ── table ──
    cols = tmpl.get("table_columns", ["No.", "Tanggal", "Uraian Aktivitas", "Paraf Penyelia"])
    # Sum must fit printable width (A4 21cm - left 4 - right 3 = 14cm).
    widths = [Cm(0.9), Cm(3.3), Cm(7.8), Cm(2.0)]
    sig_roles = tmpl.get("signature_roles", ["Peserta", "Penyelia"])

    import re as _re

    def _add_md_runs(p, text):
        """Add runs to paragraph: *...* -> italic, `...` -> monospace code."""
        for part in _re.split(r"(\*[^*]+\*|`[^`]+`)", str(text)):
            if not part:
                continue
            if len(part) >= 2 and part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1])
                run.italic = True
                run.font.name = font_name
                run.font.size = Pt(font_size)
            elif len(part) >= 2 and part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(font_size - 1)
            else:
                run = p.add_run(part)
                run.font.name = font_name
                run.font.size = Pt(font_size)

    def render_uraian(cell, value):
        """Render activity cell. List value -> bullet paragraphs; str -> single paragraph.
        Supports *...* markdown italics for technical/foreign terms."""
        if isinstance(value, list):
            for k, item in enumerate(value):
                p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                lead = p.add_run("• ")
                lead.font.name = font_name
                lead.font.size = Pt(font_size)
                _add_md_runs(p, item)
        else:
            p = cell.paragraphs[0]
            _add_md_runs(p, value)

    def add_table_and_entries(entries):
        table = doc.add_table(rows=1 + len(entries), cols=len(cols))
        table.style = "Table Grid"
        set_fixed_layout(table, widths)
        # Header row
        for i, (col, w) in enumerate(zip(cols, widths)):
            cell = table.rows[0].cells[i]
            cell.width = w
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(col)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size)
        # Data rows
        for idx, entry in enumerate(entries):
            row = table.rows[idx + 1]
            no_cell = row.cells[0]
            no_cell.width = widths[0]
            set_cell_border(no_cell)
            pn = no_cell.paragraphs[0]
            pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rn = pn.add_run(str(idx + 1))
            rn.font.name = font_name
            rn.font.size = Pt(font_size)

            tgl_cell = row.cells[1]
            tgl_cell.width = widths[1]
            set_cell_border(tgl_cell)
            pt = tgl_cell.paragraphs[0]
            rt = pt.add_run(entry.get("tanggal", ""))
            rt.font.name = font_name
            rt.font.size = Pt(font_size)

            ur_cell = row.cells[2]
            ur_cell.width = widths[2]
            set_cell_border(ur_cell)
            render_uraian(ur_cell, entry.get("items") or entry.get("uraian_aktivitas", ""))

            pf_cell = row.cells[3]
            pf_cell.width = widths[3]
            set_cell_border(pf_cell)

    def add_signature():
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        sig_names = [data.get("nama_mahasiswa", ""), data.get("nama_penyelia", "")]
        doc.add_paragraph()
        sig_table = doc.add_table(rows=3, cols=2)
        sig_rows = [
            [f"{sig_roles[0]},", f"{sig_roles[1]},"],   # role labels
            ["", ""],                                    # one tall row for signature image
            sig_names,                                   # names (bold)
        ]
        for i, row_data in enumerate(sig_rows):
            for j, text in enumerate(row_data):
                cell = sig_table.cell(i, j)
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                if i == 2:
                    run.bold = True
        # Make the middle row tall so a signature image fits without deleting rows.
        tr = sig_table.rows[1]._tr
        trPr = tr.get_or_add_trPr()
        trH = OxmlElement("w:trHeight")
        trH.set(qn("w:val"), "1700")   # ~3.0cm in twips
        trH.set(qn("w:hRule"), "atLeast")
        trPr.append(trH)
        sig_table.rows[1].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    weeks = data.get("weeks")
    if weeks:
        from docx.enum.text import WD_BREAK
        for wi, wk in enumerate(weeks):
            h = doc.add_paragraph()
            hr = h.add_run(f"{wk.get('label', '')} ({wk.get('periode', '')})")
            hr.bold = True
            hr.font.name = font_name
            hr.font.size = Pt(font_size)
            add_table_and_entries(wk.get("entries", []))
            add_signature()
            if wi < len(weeks) - 1:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    else:
        add_table_and_entries(data.get("entries", []))
        doc.add_paragraph()
        doc.add_paragraph()
        add_signature()

    output_path = versioned_path(output_path)
    doc.save(str(output_path))
    return output_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data",   required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    data   = json.loads(Path(args.data).read_text(encoding="utf-8"))
    config = load_config()
    out    = generate(data, config, Path(args.output).expanduser())
    print(json.dumps({"success": True, "output": str(out)}))


if __name__ == "__main__":
    main()
