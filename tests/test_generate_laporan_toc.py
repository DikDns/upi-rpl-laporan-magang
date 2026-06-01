# tests/test_generate_laporan_toc.py
import sys, tempfile
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent / "scripts"))
from generate_laporan import scan_has_tables, scan_has_images, _insert_toc_field

def _write(tmpdir, name, content):
    p = Path(tmpdir) / name
    p.write_text(content, encoding="utf-8")
    return p

def test_scan_has_tables_true():
    with tempfile.TemporaryDirectory() as d:
        _write(d, "bab2.md", "| Col1 | Col2 |\n|------|------|\n| A | B |")
        assert scan_has_tables(Path(d)) is True

def test_scan_has_tables_false():
    with tempfile.TemporaryDirectory() as d:
        _write(d, "bab1.md", "Tidak ada tabel di sini.")
        assert scan_has_tables(Path(d)) is False

def test_scan_has_images_true():
    with tempfile.TemporaryDirectory() as d:
        _write(d, "bab3.md", "Teks biasa\n![Gambar kerja](foto.jpg)\nTeks lagi")
        assert scan_has_images(Path(d)) is True

def test_scan_has_images_false():
    with tempfile.TemporaryDirectory() as d:
        _write(d, "bab1.md", "Tidak ada gambar di sini.")
        assert scan_has_images(Path(d)) is False

def test_insert_toc_field_adds_paragraphs():
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document()
    initial_count = len(doc.paragraphs)
    _insert_toc_field(doc, "Arial", 11, "DAFTAR ISI", 'TOC \\o "1-3" \\h \\z \\u')
    # Should add at least 2 paragraphs: title + TOC field paragraph
    assert len(doc.paragraphs) >= initial_count + 2

def test_insert_toc_field_title_text():
    from docx import Document
    doc = Document()
    _insert_toc_field(doc, "Arial", 11, "DAFTAR ISI", 'TOC \\o "1-3" \\h \\z \\u')
    texts = [p.text for p in doc.paragraphs]
    assert "DAFTAR ISI" in texts

def test_insert_toc_field_contains_fld_char():
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document()
    _insert_toc_field(doc, "Arial", 11, "DAFTAR ISI", 'TOC \\o "1-3" \\h \\z \\u')
    # Find fldChar elements in the document XML
    fld_chars = doc.element.body.findall(f".//{{{qn('w:fldChar').split('}')[0][1:]}}}fldChar") if False else \
        [el for el in doc.element.body.iter() if el.tag.endswith("}fldChar")]
    assert len(fld_chars) >= 2  # begin + end
